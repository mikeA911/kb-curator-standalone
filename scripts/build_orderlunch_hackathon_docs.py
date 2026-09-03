from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT_GUIDE = ROOT / "docs" / "showcase" / "OrderLunch-Agent-Guide.docx"
OUT_PROPOSAL = ROOT / "docs" / "hackathon" / "PSHS-Filipino-Street-Food-AI-Hackathon-Proposal.docx"

NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4D6"
GOLD = "C58A00"
GRAY = "5F6B7A"
LIGHT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_run(run, size=11, bold=False, italic=False, color="000000", font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure(doc, title, label, preset="guide"):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1.30)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6 if preset == "guide" else 8)
    normal.paragraph_format.line_spacing = 1.25 if preset == "guide" else 1.33

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14 if preset == "guide" else 12, 7 if preset == "guide" else 6),
        ("Heading 3", 12, NAVY, 10 if preset == "guide" else 8, 5 if preset == "guide" else 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.text = ""
    set_run(header.add_run(label), size=9, bold=True, color=GRAY)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    header.add_run("\t")
    page_run = header.add_run("Page ")
    set_run(page_run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    header._p.append(fld)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sec.footer.paragraphs[0].text = ""

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run(r, size=27 if preset == "guide" else 26, bold=True, color=NAVY)
    return doc


def subtitle(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    set_run(p.add_run(text), size=13, color=GRAY)


def kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text.upper()), size=9.5, bold=True, color=GOLD)


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run(p.add_run(bold_lead), bold=True, color=NAVY)
        set_run(p.add_run(text[len(bold_lead):]))
    else:
        set_run(p.add_run(text))
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text))
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(text))
    return p


def callout(doc, heading, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(heading), size=11, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(text), size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, text in enumerate(headers):
        hdr.cells[i].width = Inches(widths[i])
        set_cell_shading(hdr.cells[i], PALE_BLUE)
        set_cell_margins(hdr.cells[i])
        header_p = hdr.cells[i].paragraphs[0]
        header_p.paragraph_format.space_after = Pt(0)
        header_p.paragraph_format.line_spacing = 1.05
        set_run(header_p.add_run(text), size=9.5, bold=True, color=NAVY)
    for row in rows:
        added_row = t.add_row()
        prevent_row_split(added_row)
        cells = added_row.cells
        for i, text in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[i])
            cell_p = cells[i].paragraphs[0]
            cell_p.paragraph_format.space_after = Pt(0)
            cell_p.paragraph_format.line_spacing = 1.05
            set_run(cell_p.add_run(str(text)), size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def build_guide():
    doc = configure(Document(), "OrderLunch Agent Guide", "KB Sandbox Showcase | OrderLunch", "guide")
    kicker(doc, "Operator and integration reference")
    subtitle(doc, "How Ember—or any approved client application—can browse, quote, confirm, place, track and cancel a simulated Filipino lunch order.")
    callout(doc, "Current status", "The MCP service is live on Railway and uses persistent PostgreSQL storage. It is a safe showcase: simulated outlets only, pay upon delivery, and no real payment, address or delivery-provider integration.")

    doc.add_heading("1. What the agent demonstrates", level=1)
    para(doc, "OrderLunch demonstrates a portable enterprise agent capability. KB Sandbox guides and governs the work, Ember provides one conversational interface, and independent web, mobile, kiosk or messaging applications may call the same registered MCP service.")
    for text in (
        "A documented MCP contract rather than interface-specific business logic.",
        "Project-bound delegated identity and least-privilege tool access.",
        "Server-authoritative menu prices and availability.",
        "Explicit human confirmation before placing or cancelling an order.",
        "Idempotent placement to prevent accidental duplicate orders.",
        "Persistent status and audit records.",
        "Explicit paymentTerms: pay_on_delivery; no payment credentials are collected.",
    ): bullet(doc, text)

    doc.add_heading("2. Live service", level=1)
    table(doc, ["Item", "Value"], [
        ("MCP endpoint", "https://orderlunch-mcp-showcase-production.up.railway.app/mcp"),
        ("Health", "GET /healthz"),
        ("Readiness", "GET /readyz"),
        ("Transport", "MCP Streamable HTTP"),
        ("Authentication", "Gateway API key plus short-lived signed user delegation"),
        ("Currency", "PHP"),
        ("Payment", "Pay upon delivery"),
    ], [1.55, 4.95])
    callout(doc, "Secret handling", "Gateway keys and signing secrets belong in approved secret storage. Never place them in source code, chat messages, screenshots, ordinary database fields or documentation.", PALE_GOLD)

    doc.add_heading("3. User journey", level=1)
    for step in (
        "Start from the client-specific KB Sandbox Project or another authorized application.",
        "Ask to browse outlets or describe the desired food, quantities and fulfilment method.",
        "Review server-provided menu items, prices and availability.",
        "Request a quotation. Confirm that the items, total, fulfilment and pay-upon-delivery term are correct.",
        "Request approval. Ember must stop and display the trusted confirmation control.",
        "Confirm through the trusted interface. The model cannot confirm on the user's behalf.",
        "Place the simulated order using a server-generated idempotency key.",
        "Check status or request cancellation. Cancellation has its own confirmation gate.",
    ): numbered(doc, step)

    doc.add_heading("4. Example prompts", level=1)
    table(doc, ["Intent", "Example"], [
        ("Explore", "What lunch outlets and meals are available today?"),
        ("Quote", "Prepare a delivery quotation for two Chicken Adobo Rice meals."),
        ("Confirm terms", "Show the total and confirm whether this is pay upon delivery."),
        ("Track", "What is the status of my most recent lunch order?"),
        ("Cancel", "Please help me cancel the order. Show me what I need to confirm."),
    ], [1.3, 5.2])

    doc.add_heading("5. Available tools", level=1)
    table(doc, ["Tool", "Purpose", "Control"], [
        ("list_outlets", "List simulated food outlets", "Read only"),
        ("browse_menu", "Return menu and authoritative prices", "Read only"),
        ("check_availability", "Check item availability", "Read only"),
        ("prepare_quotation", "Create expiring immutable quotation", "No purchase"),
        ("request_order_approval", "Create pending approval", "Human gate follows"),
        ("place_order", "Place confirmed simulated order", "Consequential + idempotent"),
        ("get_order_status", "Read caller's order", "User/Project scoped"),
        ("cancel_order", "Cancel eligible order", "Separate confirmation"),
        ("advance_order_state", "Advance demonstration state", "Test operator only"),
    ], [1.65, 3.0, 1.85])

    doc.add_heading("6. Delivery address design", level=1)
    para(doc, "The current showcase does not collect an address. The recommended next phase keeps personal addresses outside model prompts and ordinary chat history:")
    for text in (
        "Ember opens a trusted address form after the user selects Delivery.",
        "The user selects an approved office location, a private saved address or a one-time address.",
        "The trusted service validates the address and returns an opaque deliveryAddressId.",
        "The MCP order receives the identifier—not the full address in the prompt.",
        "The final confirmation shows the human-readable address, order total and pay-upon-delivery terms.",
        "Addresses are encrypted, access-controlled, excluded from logs, and shared with a delivery provider only after confirmation.",
    ): bullet(doc, text)

    doc.add_heading("7. Images and storefront presentation", level=1)
    para(doc, "A future version can return hosted storefront and menu-image URLs together with structured item data. Ember and independent applications can render visual menu cards while retaining server-authoritative names, prices, allergens and availability. The model must never invent images, prices or availability.")

    doc.add_heading("8. Safety checklist", level=1)
    for text in (
        "Bind each call to the authenticated user and Project.",
        "Grant only required tools; reserve advance_order_state for demo operators.",
        "Never expose the human-confirmation operation as an MCP tool.",
        "Show the exact basket, price, payment term and expiry before confirmation.",
        "Keep personal addresses and secrets out of model context and logs.",
        "Treat allergens, food safety, merchant onboarding and real delivery as separate governed capabilities.",
    ): bullet(doc, text)

    doc.add_heading("9. Current limitations", level=1)
    para(doc, "The live service uses simulated outlets and menus. It does not contact merchants, reserve inventory, collect payment, store addresses, dispatch riders, calculate taxes or delivery fees, or guarantee food safety. Those capabilities require new contracts, evidence, controls and operator agreements before any real-world use.")
    OUT_GUIDE.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_GUIDE)


def build_proposal():
    doc = configure(Document(), "Philippine Science High School Filipino Street-Food AI Hackathon", "Draft Proposal | For Discussion", "proposal")
    kicker(doc, "A Philippine builder ecosystem for trusted, deployable AI")
    subtitle(doc, "Students design, govern, build and demonstrate a portable fishball-stand agent—then connect it to Ember or an application of their own.")
    callout(doc, "Proposal at a glance", "A supervised student hackathon using KB Sandbox for evidence-led requirements, architecture, safety and evaluation; assisted coding tools for implementation; Sandz infrastructure for deployment; and a prominent fishball-stand showcase near Sandz in the Ayala area of Makati as the proposed winning deployment opportunity.")

    doc.add_heading("1. Purpose", level=1)
    para(doc, "The hackathon gives Philippine Science High School students practical experience with the skills becoming central to AI-enabled software development: precise requirements, evidence gathering, architecture, tool contracts, privacy, human approval, testing and critical evaluation—not merely generating code.")
    para(doc, "The challenge is intentionally familiar and local. Teams make a Filipino street-food outlet AI-ready, beginning with fishballs and potentially extending to kwek-kwek, squid balls, kikiam, sauces and drinks. The resulting service must be portable: Ember is one client, while student-built web, mobile, kiosk or messaging applications may call the same agent or MCP server.")

    doc.add_heading("2. Proposed organizers and beneficiaries", level=1)
    table(doc, ["Participant", "Proposed role"], [
        ("Philippine Science High School", "Student participation, faculty supervision, safeguarding and academic alignment"),
        ("Sandz", "Programme sponsor, infrastructure partner, mentor access and showcase location coordination"),
        ("KB Sandbox", "Foundation knowledge, Methods, Project workspace, evidence governance, evaluation and Ember demonstration interface"),
        ("Local software builders", "Mentoring, API/MCP review, implementation coaching and potential apprenticeship pathways"),
        ("Outlet/operator partner", "Menu and operating input, food-safety ownership, acceptance testing and real-world feedback"),
    ], [2.0, 4.5])
    callout(doc, "Approval status", "This is a concept proposal, not a confirmed school programme or prize commitment. Participation, use of school name, treatment of minors, venue, stand location, food operations and awards require written approval from the relevant organizations.", PALE_GOLD)

    doc.add_heading("3. Challenge statement", level=1)
    para(doc, "Design and build a safe, usable and portable AI ordering capability for a Filipino street-food stand. The solution should help a customer discover the menu, receive a reliable quotation, confirm pay-upon-delivery terms, place an order only after explicit approval, and track or cancel the order safely.")
    para(doc, "Teams may build their own customer-facing application. They must also register or connect the agent/MCP service to a KB Sandbox Project so evaluators can review its evidence, architecture, controls and behavior through Ember.")

    doc.add_heading("4. Required team deliverables", level=1)
    for text in (
        "Problem and user-story definition, including customer and stand-operator journeys.",
        "Evidence pack: menu, prices, ingredients/allergens, operating assumptions, privacy considerations and constraints.",
        "OpenAPI-equivalent capability contract and/or MCP tool specification.",
        "Architecture and threat model showing identity, permissions, data movement, human gates, idempotency and audit.",
        "Working agent or MCP server deployed on approved Sandz-hosted infrastructure.",
        "At least one client interface: Ember integration, student-built app, kiosk or another approved interface.",
        "Automated and human test evidence, including failures and limitations.",
        "Short demonstration, operator handover and roadmap for real-world readiness.",
    ): bullet(doc, text)

    doc.add_heading("5. Minimum functional scope", level=1)
    table(doc, ["Capability", "Minimum expectation"], [
        ("Menu", "Structured item names, images where licensed, authoritative prices, availability and allergen notices"),
        ("Quotation", "Exact basket, PHP total, fulfilment choice, expiry and pay-upon-delivery term"),
        ("Confirmation", "Trusted human approval separate from the model/agent"),
        ("Order", "Idempotent placement and persistent order identifier"),
        ("Status", "Clear simulated lifecycle such as placed, preparing, ready, completed or cancelled"),
        ("Cancellation", "Separate confirmation and appropriate state restrictions"),
        ("Identity", "User- and Project-scoped access; no identity supplied by prompt text"),
        ("Privacy", "No secret or full delivery address in prompts, logs or public demonstrations"),
        ("Portability", "Documented contract usable by Ember and at least one independent client"),
    ], [1.45, 5.05])

    doc.add_heading("6. Suggested programme stages", level=1)
    stages = [
        ("1. Foundation", "Learn AI, LLM, RAG, agents, MCP, APIs, evidence, privacy, human approval and evaluation."),
        ("2. Qualification", "Complete an evidence-grounded assessment using the KB Sandbox foundation knowledge base."),
        ("3. Project design", "Create a team Project, gather evidence and produce the required architecture deliverables."),
        ("4. Build sprint", "Implement the server and client with approved assisted-coding tools and mentor checkpoints."),
        ("5. Register and test", "Connect the service to KB Sandbox, certify tool behavior and run repeated evaluations."),
        ("6. Demonstration", "Present the customer journey, operator view, architecture, safety controls, evidence and lessons."),
        ("7. Showcase", "Subject to approvals, harden and deploy the winning solution for a demonstration fishball stand."),
    ]
    table(doc, ["Stage", "Outcome"], stages, [1.45, 5.05])

    doc.add_heading("7. Judging framework", level=1)
    table(doc, ["Criterion", "Points", "What judges examine"], [
        ("Customer and operator value", "15", "Clarity, usefulness, cultural fit and ease of operation"),
        ("Requirements and evidence", "15", "Traceability, assumptions, source quality and honest gaps"),
        ("Architecture and portability", "15", "Contract quality, client independence and maintainable boundaries"),
        ("Safety, privacy and governance", "20", "Identity, permissions, confirmation, address handling, allergens and audit"),
        ("Reliability and testing", "15", "Idempotency, failure handling, persistence, repeatability and test evidence"),
        ("User experience", "10", "Accessible interface, clear totals/terms and understandable errors"),
        ("Communication and learning", "10", "Demonstration, documentation, reflection and response to questions"),
    ], [2.0, 0.65, 3.85])

    doc.add_heading("8. Proposed prize and recognition", level=1)
    para(doc, "The proposed headline opportunity is to prepare the winning solution for a visible demonstration at a fishball stand near Sandz in the Ayala area of Makati. The final arrangement may be a sponsored showcase, pilot deployment or equivalent prize depending on location, permits, operator agreement, food-safety requirements and school approval.")
    for text in (
        "Winning team recognition and certificates for finalists.",
        "Mentoring to harden the winning prototype into a governed showcase.",
        "Sandz-hosted deployment for an agreed demonstration period.",
        "Portfolio-ready architecture, evaluation and deployment evidence for students.",
        "Potential internship, apprenticeship or builder-community introductions subject to separate selection processes.",
    ): bullet(doc, text)

    doc.add_heading("9. Safeguarding and responsible participation", level=1)
    for text in (
        "Obtain school authorization, parental/guardian consent where required, and appropriate supervision for minors.",
        "Do not use real customer addresses, payment information, private accounts or production merchant credentials during competition work.",
        "Use simulated orders until the organizers approve a separately governed real-world pilot.",
        "Require licensed or team-created images and documented source attribution.",
        "Include allergen notices and make clear that the software does not replace food-safety responsibility.",
        "Prohibit harassment, unsafe content, credential sharing, unauthorized scanning and use of third-party systems without permission.",
        "Provide an incident-reporting route and an adult safeguarding contact throughout the event.",
    ): bullet(doc, text)

    doc.add_heading("10. Intellectual property and portability", level=1)
    para(doc, "Before launch, organizers should publish simple terms explaining ownership, licensing, open-source obligations, use of student work in demonstrations, attribution, and any rights needed to operate the winning showcase. Students should retain appropriate recognition. Any commercialization, revenue sharing, maintenance obligation or transfer of ownership should require a separate written agreement rather than being implied by participation.")

    success_heading = doc.add_heading("11. Proposed success measures", level=1)
    success_heading.paragraph_format.page_break_before = True
    table(doc, ["Area", "Suggested measure"], [
        ("Learning", "Improvement between foundation assessment and final architecture/evaluation review"),
        ("Completion", "Teams delivering all required evidence, contract, build and test artifacts"),
        ("Reliability", "Successful repeated order flows with no duplicate placement or cross-user leakage"),
        ("Portability", "Working through Ember and an independent client application"),
        ("Safety", "No exposed secrets or personal information; all consequential actions retain human gates"),
        ("Builder ecosystem", "Students, mentors and local software houses continuing with further governed projects"),
        ("Showcase readiness", "One solution suitable for a supervised public demonstration after hardening"),
    ], [1.55, 4.95])

    doc.add_heading("12. Decisions required before announcement", level=1)
    for text in (
        "Confirm participating PSHS campus/campuses, faculty sponsor and student eligibility.",
        "Confirm dates, team size, format, venue, mentor pool and safeguarding arrangements.",
        "Confirm Sandz sponsorship, hosting limits and demonstration-site feasibility.",
        "Approve challenge rules, judging panel, IP/license terms and acceptable-use policy.",
        "Define whether the final showcase remains simulated or progresses to a separately approved real outlet pilot.",
        "Prepare the KB Sandbox foundation assessment, Project template, Methods and evaluator rubric.",
    ): numbered(doc, text)

    doc.add_heading("Appendix A. Reference architecture", level=1)
    para(doc, "Customer interface (Ember, student app or kiosk) -> KB Sandbox Project and/or authorized client gateway -> signed user delegation -> Fishball Stand agent/MCP server -> menu and order state -> trusted human confirmation -> simulated operator workflow.")
    para(doc, "KB Sandbox remains the evidence, architecture, governance and evaluation layer. The agent remains portable and may serve other authorized applications through its documented contract.")

    doc.add_heading("Appendix B. Expansion ideas", level=1)
    for text in (
        "Tagalog and regional-language ordering assistance.",
        "Visual menus, dietary filters and allergen warnings.",
        "Secure one-time delivery addresses represented to the agent by opaque identifiers.",
        "Stand inventory and low-stock notifications.",
        "Delivery-provider integration through a separately governed connector or MCP server.",
        "Invoice creation and processing for larger outlet or catering orders.",
        "Reusable AI-readiness packages for other microbusinesses and community services.",
    ): bullet(doc, text)

    OUT_PROPOSAL.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PROPOSAL)


if __name__ == "__main__":
    build_guide()
    build_proposal()
    print(OUT_GUIDE)
    print(OUT_PROPOSAL)
